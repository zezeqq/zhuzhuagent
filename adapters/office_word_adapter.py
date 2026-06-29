from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

from artifacts.artifact_manager import ensure_export_dir


def normalize_word_sections(
    sections: list[Any] | None = None,
    content: str | None = None,
) -> list[tuple[str, str]]:
    fixed: list[tuple[str, str]] = []
    for item in sections or []:
        if isinstance(item, dict):
            heading = str(item.get("heading") or item.get("title") or "章节")
            body = str(item.get("body") or item.get("content") or "")
            fixed.append((heading, body))
        elif isinstance(item, (list, tuple)) and len(item) >= 2:
            fixed.append((str(item[0]), str(item[1])))
    if fixed:
        return fixed
    text = (content or "").strip()
    if text:
        return [("正文", text)]
    return []


def word_input_has_data(
    sections: list[Any] | None = None,
    content: str | None = None,
) -> bool:
    return bool(normalize_word_sections(sections=sections, content=content))


def validate_word_input(
    sections: list[Any] | None = None,
    content: str | None = None,
) -> None:
    if word_input_has_data(sections, content):
        return
    raise ValueError(
        "Word 缺少有效内容：请提供 sections（heading/body）或 content 正文。"
    )


def create_word_document(
    title: str = "",
    sections: list[Any] | None = None,
    output_name: str = "document.docx",
    *,
    content: str | None = None,
) -> Path:
    normalized = normalize_word_sections(sections=sections, content=content)
    validate_word_input(sections=normalized, content=None)

    out_dir = ensure_export_dir()
    path = out_dir / output_name
    document = Document()
    document.add_heading(title or "文档", level=0)
    for heading, body in normalized:
        document.add_heading(heading, level=1)
        for paragraph in body.splitlines():
            paragraph = paragraph.strip()
            if paragraph:
                document.add_paragraph(paragraph)
    document.save(path)
    return path
