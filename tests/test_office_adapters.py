"""Office 适配器：真实落盘，不依赖 LLM。"""

from __future__ import annotations

from pathlib import Path

from adapters.office_excel_adapter import create_excel_workbook
from adapters.office_ppt_adapter import create_presentation
from adapters.office_word_adapter import create_word_document
from docx import Document
from openpyxl import load_workbook
from pptx import Presentation


def test_create_word_document(app_tmp):
    path = create_word_document(
        "测试文档",
        [("概述", "第一段\n第二段")],
        "unit_test.docx",
    )
    assert path.exists()
    doc = Document(path)
    assert doc.paragraphs[0].text == "测试文档"


def test_create_word_document_with_content(app_tmp):
    path = create_word_document(
        "产品需求",
        output_name="unit_test_content.docx",
        content="第一章 背景\n\n这是正文内容。",
    )
    assert path.exists()
    doc = Document(path)
    assert "正文内容" in doc.paragraphs[-1].text


def test_create_excel_workbook(app_tmp):
    path = create_excel_workbook(
        "清单",
        ["序号", "名称"],
        [[1, "A"], [2, "B"]],
        "unit_test.xlsx",
    )
    assert path.exists()
    wb = load_workbook(path)
    assert wb.active["A1"].value == "序号"
    assert wb.active["A2"].value == 1


def test_create_excel_workbook_with_sheets(app_tmp):
    path = create_excel_workbook(
        output_name="unit_test_sheets.xlsx",
        sheets=[
            {"title": "清单", "headers": ["序号", "名称"], "rows": [[1, "A"]]},
            {"title": "汇总", "headers": ["类别", "数量"], "rows": [["资料", 3]]},
        ],
    )
    assert path.exists()
    wb = load_workbook(path)
    assert len(wb.sheetnames) == 2
    assert wb["清单"]["A1"].value == "序号"
    assert wb["汇总"]["A1"].value == "类别"


def test_create_excel_workbook_rejects_empty(app_tmp):
    import pytest
    from adapters.office_excel_adapter import validate_excel_input

    with pytest.raises(ValueError, match="缺少有效数据"):
        validate_excel_input(title="空白表")


def test_create_presentation(app_tmp):
    path = create_presentation(
        "汇报",
        [("封面", ["要点一", "要点二"])],
        "unit_test.pptx",
    )
    assert path.exists()
    prs = Presentation(path)
    assert len(prs.slides) >= 1
